"""
resume_parser.py — Resume Upload & Skill Extraction
=====================================================
Supports PDF, DOCX, and TXT file uploads.
Extracts skills using Python-based text extraction and keyword matching.
No external AI APIs are used.
"""

import re
import io

# ─── Attempt imports (graceful handling) ──────────────────────────────────────
try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─── Comprehensive Skill Keywords ────────────────────────────────────────────
# Organized for efficient matching against resume text
RESUME_SKILL_KEYWORDS = {
    # Programming Languages
    "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "R",
    "Go", "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "MATLAB",
    "SQL", "Bash", "Perl", "Solidity",

    # Data Science & ML
    "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "PyTorch", "Keras",
    "Matplotlib", "Seaborn", "Plotly", "SciPy", "XGBoost", "LightGBM",
    "OpenCV", "spaCy", "NLTK", "Hugging Face", "Deep Learning",
    "Machine Learning", "Neural Networks", "NLP", "Computer Vision",
    "Data Wrangling", "Feature Engineering", "MLOps",

    # Web Development
    "HTML", "CSS", "React", "Angular", "Vue.js", "Node.js", "Express",
    "Django", "Flask", "FastAPI", "Next.js", "Tailwind", "Bootstrap",
    "jQuery", "GraphQL", "REST APIs", "WebSockets",

    # Databases
    "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch",
    "Firebase", "SQLite", "Oracle", "DynamoDB", "Cassandra",

    # Cloud & DevOps
    "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform",
    "Ansible", "Jenkins", "CI/CD", "Linux", "Nginx",
    "CloudFormation", "Prometheus", "Grafana",

    # Data Engineering
    "Apache Spark", "Airflow", "Kafka", "Hadoop", "ETL",
    "Data Modeling", "Snowflake", "BigQuery", "DBT",

    # Tools
    "Git", "GitHub", "GitLab", "Jira", "Confluence", "Agile",
    "Scrum", "Figma", "Adobe XD", "Tableau", "Power BI", "Excel",
    "Jupyter Notebooks", "VS Code",

    # Security
    "Network Security", "Penetration Testing", "SIEM", "Wireshark",
    "OWASP", "Risk Assessment", "Firewalls",

    # Concepts
    "Data Structures", "Algorithms", "System Design", "Microservices",
    "Design Patterns", "Object Oriented Programming", "Unit Testing",
    "Responsive Design", "Accessibility",
}


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    if not HAS_PYPDF2:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error reading PDF: {e}"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    if not HAS_DOCX:
        return ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        return f"Error reading DOCX: {e}"


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return file_bytes.decode("latin-1", errors="ignore")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Route to the appropriate extractor based on file extension.
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        return ""


def extract_skills_from_resume(file_bytes: bytes, filename: str) -> list[str]:
    """
    Extract skills from an uploaded resume file.

    Pipeline:
    1. Extract raw text from the file.
    2. Normalize the text.
    3. Match against the skill keyword database.
    4. Return a deduplicated, sorted list of found skills.

    Parameters
    ----------
    file_bytes : bytes
        Raw file content.
    filename : str
        Original filename (used to determine format).

    Returns
    -------
    list[str]
        Sorted list of extracted skills.
    """
    text = extract_text(file_bytes, filename)
    if not text or text.startswith("Error"):
        return []

    text_lower = text.lower()
    found_skills = set()

    for skill in RESUME_SKILL_KEYWORDS:
        skill_lower = skill.lower()
        # Use word boundary matching for accuracy
        # Escape special regex characters (e.g., C++, C#)
        pattern = r'\b' + re.escape(skill_lower) + r'\b'

        # Special handling for very short terms to avoid false positives
        if len(skill) <= 2:
            # For "R", "C", "Go" — require specific context patterns
            if skill == "R":
                if re.search(r'\bR\b', text):  # case-sensitive for R
                    found_skills.add(skill)
            elif skill == "Go":
                if re.search(r'\bGo\b', text) or re.search(r'\bgolang\b', text_lower):
                    found_skills.add(skill)
            elif skill == "C":
                if re.search(r'\bC\b', text):
                    found_skills.add(skill)
        else:
            if re.search(pattern, text_lower):
                found_skills.add(skill)

    return sorted(found_skills)


def get_resume_summary(text: str) -> dict:
    """
    Generate a brief summary of the resume content.
    """
    word_count = len(text.split())
    line_count = len(text.strip().split("\n"))

    # Try to extract name (usually the first non-empty line)
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    probable_name = lines[0] if lines else "Unknown"

    # Try to detect email
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
    email = email_match.group(0) if email_match else None

    # Try to detect phone
    phone_match = re.search(r'[\+]?[\d\s\-\(\)]{10,15}', text)
    phone = phone_match.group(0).strip() if phone_match else None

    return {
        "word_count": word_count,
        "line_count": line_count,
        "probable_name": probable_name,
        "email": email,
        "phone": phone,
    }
